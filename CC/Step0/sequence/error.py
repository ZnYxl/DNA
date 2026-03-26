# -*- coding: utf-8 -*-
import sys
sys.path.append('/mnt/st_data/linhaiyan/python_pkgs')
import numpy as np

from datetime import datetime


# 首先定义类

class Error_stats:
    def __init__(
            self,
            ref_path,
            read_path,
            name='error_count_file_ID20',
            output_mode=0
    ):

        # 接下来生成原始序列的dict，key为序列index，value为序列内容
        ref_data = open(ref_path, "r", encoding='utf-8')
        self.ref_dict = {}
        for line in ref_data:
            # 去除末尾换行符
            line = line.replace('\n', '')
            if line.startswith('>'):
                # 去除 > 号
                seq_name = line[1:]
                self.ref_dict[seq_name] = ''
            else:
                # 去除末尾换行符并连接多行序列
                self.ref_dict[seq_name] += line.replace('\n', '')

        self.read_data = open(read_path, "r", encoding='utf-8')
        self.output_mode = output_mode
        now = datetime.now()
        time_str = now.strftime("%Y-%m-%H-%M")
        if name == '':
            self.name = time_str
        else:
            self.name = name + "-" + time_str

        self.error_dict = {
            "total_reads": 0,  # 总序列数
            "ref_len": 0,  # 原始序列长度
            "total_ref": 0,  # 总原始序列数
            "read_length_list": [0] * 300,  # 序列长度分布表
            "total_base_nums": 0,  # 总碱基数
            "total_error_nums": 0,  # 总错误碱基数
            "total_del_error_nums": 0,  # 总删除碱基数
            "total_sub_error_nums": 0,  # 总替换错误碱基数
            "total_ins_error_nums": 0,  # 总增添碱基数
            "total_error_reads_nums": 0,  # 总错误序列数
            "total_broken_reads_nums": 0,  # 总断链序列数
            "total_no_align_nums": 0,  # 未比对上序列数

            "miss_ref_nums": 0,  # 缺失原始序列数量
            "miss_ref_list": {},  # 原始序列比对数分布字典
            "miss_list": [],

            # 接下来是不同碱基对于错误率的影响：
            "base_del_nums": [0, 0, 0, 0],  # 不同碱基的删除数，依次是ATGC
            "base_ins_nums": [0, 0, 0, 0],  # 不同碱基的增添数，依次是ATGC
            "base_sub_nums_matrix": np.zeros((4, 4)),  # 依次是ATGC的替换矩阵，其中第一维为原来碱基，第二维为替换后的碱基

            # 接下来是碱基位置的影响
            "loc_base_error_list": [0] * 300,  # 不同位置的碱基错误率
            "loc_base_del_error_list": [0] * 300,  # 不同位置的碱基删除率
            "loc_base_sub_error_list": [0] * 300,  # 不同位置的碱基替换率
            "loc_base_ins_error_list": [0] * 300,  # 不同位置的碱基增添率

            # 接下来是碱基连续性的影响：
            "pre_order_del_matrix": np.zeros((4, 4)),  # 前序碱基对于后续删除错误率的影响矩阵，第一维为前序碱基，第二维为后序碱基
            "pre_order_sub_matrix": np.zeros((4, 4)),  # 前序碱基对于后续替换错误率的影响矩阵，第一维为前序碱基，第二维为后序碱基
            "pre_order_ins_matrix": np.zeros((4, 4)),  # 前序碱基对于后续增添错误率的影响矩阵，第一维为前序碱基，第二维为后序碱基
            "post_order_error_matrix": np.zeros((4, 4)),  # 后序碱基对于后续错误率的影响矩阵，第一维为后序碱基，第二维为前序碱基
            "post_order_sub_error_matrix": np.zeros((4, 4)),  # 后序碱基对于后续替换错误率的影响矩阵，第一维为后序碱基，第二维为前序碱基
            "post_order_ins_error_matrix": np.zeros((4, 4)),  # 后序碱基对于后续增添错误率的影响矩阵，第一维为后序碱基，第二维为前序碱基

            # 接下来是连续性错误的情况：
            "contiguity_error_list": [0] * 300,  # 连续多个错误的长度分布情况
            "contiguity_del_error_list": [0] * 300,  # 连续多个删除错误的长度分布情况
            "contiguity_sub_error_list": [0] * 300,  # 连续多个替换错误的长度分布情况
            "contiguity_ins_error_list": [0] * 300,  # 连续多个增添错误的长度分布情况
            "broken_error_list": [0] * 300,  # 断链长度分布情况
            "broken_error_base_list": np.zeros((4, 4))  # 断链的位置碱基情况, 第一维为断链前碱基，第二维为断链后碱基
        }

        # 接下来把统计字典的一些参数给初始化一下
        first_key = next(iter(self.ref_dict))
        first_value = self.ref_dict[first_key]
        self.error_dict["ref_len"] = len(first_value)
        self.error_dict["total_ref"] = len(self.ref_dict)
        self.error_dict["total_base_nums"] = self.error_dict["ref_len"] * self.error_dict["total_ref"]
        for key in self.ref_dict:
            self.error_dict["miss_ref_list"][key] = 0

    def align_uint_fuc(self, ref, read):
        """
        该函数为两条序列比对的单元函数，该函数输入两条序列，输出比对结果并将统计结果输入到self.error_dict
        """
        # 注意 ref为ref，read为测序后read
        len_ref = len(ref)
        len_read = len(read)
        base2nums_dict = {"A": 0, "T": 1, "G": 2, "C": 3}

        # print(ref,read)
        # 初始化二维数组
        dp = [[0 for _ in range(len_read + 1)] for _ in range(len_ref + 1)]

        for i in range(len_ref + 1):
            dp[i][0] = i

        for j in range(len_read + 1):
            dp[0][j] = j

        for i in range(1, len_ref + 1):
            for j in range(1, len_read + 1):
                if ref[i - 1] == read[j - 1]:
                    dp[i][j] = dp[i - 1][j - 1]
                else:
                    dp[i][j] = min(dp[i - 1][j - 1] + 1, dp[i - 1][j] + 1, dp[i][j - 1] + 1)

        # 回溯找到替换位置

        # 我们将错误放到一个统计list里面，元素的结构为（index,kind,（a,b,c）,err）
        # 其中index为错误发生的位置（以ref为准），kind为发生的种类，0为删除，1为替换，2为增添，3为断链  （abc）为ref中前后中的碱基，err为错误碱基

        error_uint_list = []
        i, j = len_ref, len_read
        operations = []

        # 先检验下断链的情况
        broken_length = 0
        broken_base = 0
        while ref[i - 1] != read[j - 1] and dp[i][j] == dp[i - 1][j] + 1:
            broken_length += 1
            broken_base = (ref[i - 2], ref[i - 1])
            i -= 1

        while i > 0 and j > 0:
            if ref[i - 1] == read[j - 1]:
                i, j = i - 1, j - 1
            elif dp[i][j] == dp[i - 1][j - 1] + 1:
                # 表示发生了替换错误

                error_uint_list.append((i, 1, (ref[max(0, i - 2)], ref[i - 1], ref[min(i, len_ref - 1)]), read[j - 1]))
                # operations.append(f"Replace '{ref[i - 1]}' at position {i} with '{read[j - 1]}'")
                i, j = i - 1, j - 1
            elif dp[i][j] == dp[i - 1][j] + 1:
                # 表示发生了删除错误
                error_uint_list.append((i, 0, (ref[max(0, i - 2)], ref[i - 1], ref[min(i, len_ref - 1)]), "A"))
                # operations.append(f"Delete '{ref[i - 1]}' at position {i}")
                i -= 1
            elif dp[i][j] == dp[i][j - 1] + 1:
                # 表示发生增添错误
                # print(i,j)
                error_uint_list.append((i + 1, 2,
                                        (ref[max(0, i - 1)], ref[min(i, len_ref - 1)], ref[min(i + 1, len_ref - 1)]),
                                        read[j - 1]))
                # operations.append(f"Insert '{read[j - 1]}' at position {i + 1}")
                j -= 1

        while i > 0:
            # 前面的断链
            error_uint_list.append((i, 4, (ref[max(0, i - 2)], ref[i - 1], ref[min(i, len_ref - 1)]), "A"))
            i -= 1
            # operations.append(f"Delete '{ref[i - 1]}' at position {i}")

        while j > 0:
            # 前面的增添链
            error_uint_list.append((1, 2, (ref[i - 1], ref[i], ref[min(i + 1, len_ref - 1)]), read[j - 1]))
            operations.append(f"Insert '{read[j - 1]}' at position 1")
            j -= 1

        # error_uint_list.reverse()

        # 接下来统计序列的错误信息

        self.error_dict["read_length_list"][len(read)] += 1

        if broken_length > 0:
            # 记录一下断链情况
            self.error_dict["broken_error_list"][broken_length] += 1
            self.error_dict["total_broken_reads_nums"] += 1
            self.error_dict["broken_error_base_list"][
                base2nums_dict[broken_base[0]], base2nums_dict[broken_base[1]]] += 1

        if error_uint_list != []:
            # print(error_uint_list)
            # 记录一下错链
            self.error_dict["total_error_reads_nums"] += 1
            last_kind = 5
            last_index = 0
            contiguity_length = 1
            if "N" not in read:

                for error_inidex, line in enumerate(error_uint_list):
                    index = int(line[0])
                    kind = int(line[1])
                    ref_base_list = line[2]
                    error_base = line[3]
                    error_base_index = base2nums_dict[error_base]
                    pre_error_index = base2nums_dict[ref_base_list[0]]
                    error_ref_index = base2nums_dict[ref_base_list[1]]
                    post_error_index = base2nums_dict[ref_base_list[2]]
                    # 依据错误类别分类统计
                    if kind == 0:

                        # 删除错误
                        self.error_dict["total_del_error_nums"] += 1
                        # 删除碱基统计
                        self.error_dict["base_del_nums"][error_ref_index] += 1
                        # 删除位置统计
                        self.error_dict["loc_base_del_error_list"][index] += 1
                        # 前序对后序删除的影响统计
                        self.error_dict["pre_order_del_matrix"][pre_error_index, error_ref_index] += 1
                        # 连续性错误的统计
                        if last_index - index == 1 and last_kind == kind:
                            last_index = index
                            contiguity_length += 1
                        else:
                            if contiguity_length > 1:
                                self.error_dict["contiguity_del_error_list"][contiguity_length] += 1
                            last_index = index
                            last_kind = kind
                            contiguity_length = 1
                        if error_inidex == len(error_uint_list) - 1 and contiguity_length > 1:
                            self.error_dict["contiguity_del_error_list"][contiguity_length] += 1
                    elif kind == 1:
                        # 替换错误
                        self.error_dict["total_sub_error_nums"] += 1
                        # 替换碱基统计
                        self.error_dict["base_sub_nums_matrix"][error_ref_index, error_base_index] += 1
                        # 替换位置统计
                        self.error_dict["loc_base_sub_error_list"][index] += 1
                        # 前序对后序替换的影响统计
                        self.error_dict["pre_order_sub_matrix"][pre_error_index, error_base_index] += 1
                        # 后序碱基对于替换的影响统计
                        self.error_dict["post_order_sub_error_matrix"][post_error_index, error_base_index] += 1
                        # 连续性错误的统计
                        if last_index - index == 1 and last_kind == kind:
                            last_index = index
                            contiguity_length += 1
                        else:
                            if contiguity_length > 1:
                                self.error_dict["contiguity_sub_error_list"][contiguity_length] += 1
                            last_index = index
                            last_kind = kind
                            contiguity_length = 1
                        if error_inidex == len(error_uint_list) - 1 and contiguity_length > 1:
                            self.error_dict["contiguity_sub_error_list"][contiguity_length] += 1

                    elif kind == 2:

                        # 增添错误
                        self.error_dict["total_ins_error_nums"] += 1
                        # 增添碱基统计
                        self.error_dict["base_ins_nums"][error_base_index] += 1
                        # 增添位置统计
                        self.error_dict["loc_base_ins_error_list"][index] += 1
                        # 前序对后序增添的影响统计
                        self.error_dict["pre_order_ins_matrix"][pre_error_index, error_base_index] += 1
                        # 后序碱基对于增添的影响统计
                        self.error_dict["post_order_ins_error_matrix"][post_error_index, error_base_index] += 1
                        # 连续性错误的统计
                        if last_index - index == 0 and last_kind == kind:
                            last_index = index
                            contiguity_length += 1
                        else:
                            if contiguity_length > 1:
                                self.error_dict["contiguity_ins_error_list"][contiguity_length] += 1
                            last_index = index
                            last_kind = kind
                            contiguity_length = 1
                        if error_inidex == len(error_uint_list) - 1 and contiguity_length > 1:
                            self.error_dict["contiguity_ins_error_list"][contiguity_length] += 1

    # 现在定义一个主函数，负责对输入的SAM文件进行错误统计

    def process_read(self):
        i = 0
        while True:
            read = self.read_data.readline()
            if "@" in read:
                continue
            if read == [] or read == "":
                break
            i += 1
            if i % 10000 == 0:
                print("已处理序列数：", i)
                output_name = self.name + "_some_error_dict"
                f = open(output_name, "w")
                f.write(str(self.error_dict))
                f.close()
            # print(read)
            info_list = read.split()
            # print(info_list)
            flag = info_list[1]
            ref_index = info_list[2]
            read_info = info_list[9]

            self.error_dict["total_reads"] += 1

            if ref_index == "*":
                self.error_dict["total_no_align_nums"] += 1
                continue
            else:
                self.error_dict["miss_ref_list"][ref_index] += 1
                ref_info = self.ref_dict[ref_index]
                self.align_uint_fuc(ref_info, read_info)

        # 接下来再整理下错误dict
        self.error_dict["total_error_nums"] = self.error_dict["total_del_error_nums"] + self.error_dict[
            "total_sub_error_nums"] + self.error_dict["total_ins_error_nums"]

        for key in self.error_dict["miss_ref_list"]:
            if self.error_dict["miss_ref_list"][key] == 0:
                self.error_dict["miss_ref_nums"] += 1
                self.error_dict["miss_list"].append(key)
        self.error_dict["loc_base_error_list"] = [sum(x) for x in zip(self.error_dict["loc_base_del_error_list"],
                                                                      self.error_dict["loc_base_sub_error_list"],
                                                                      self.error_dict["loc_base_ins_error_list"])]
        self.error_dict["contiguity_error_list"] = [sum(x) for x in zip(self.error_dict["contiguity_del_error_list"],
                                                                        self.error_dict["contiguity_sub_error_list"],
                                                                        self.error_dict["contiguity_ins_error_list"])]

        # 把dict写出来
        output_name = self.name + "_error_dict"
        f = open(output_name, "w")
        f.write(str(self.error_dict))
        f.close()
        # 接下来写到word里面
        self.output_word()



    def output_word(self, path=""):
        if path != "":
            # g= open(path,"r").read()
            self.error_dict = path

        from docx import Document
        from docx.shared import Inches, Pt
        from PIL import Image
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.shared import RGBColor
        # 创建一个新的Word文档
        doc = Document()
        doc.encoding = 'utf-8'
        doc.styles['Normal'].font.name = u'Arial'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        text = self.name + "错误分析\n"
        p = doc.add_paragraph(text)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.runs[0]
        font = run.font
        font.name = '微软雅黑'
        font.size = Pt(25)  # 字体大小

        title_text = "总体序列情况"
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(8, 46, 84)

        # 添加序列信息
        # txt ="总序列数："+ "\n" +"原始序列数" + +"原始序列长度" + "序列长度分布表"
        self.get_txt(doc, "总序列数：" + str(self.error_dict["total_reads"]))
        self.get_txt(doc, "原始序列数：" + str(self.error_dict["total_ref"]))
        self.get_txt(doc, "原始序列长度：" + str(self.error_dict["ref_len"]))
        self.get_txt(doc, "序列长度分布表：")
        self.get_fig(doc, "序列长度", self.error_dict["read_length_list"])

        title_text = "总体错误情况"
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(8, 46, 84)

        self.error_dict["total_base_nums"] = (self.error_dict["total_reads"] - self.error_dict["total_no_align_nums"]) * \
                                             self.error_dict["ref_len"]

        # 添加总体错误信息
        txt2 = "总体错误率：" + "总体删除率" + "总体替换率" + "总体增添率" + "总错误序列率" + "总断链率"
        self.get_txt(doc,
                     "总体错误率：" + str((self.error_dict["total_error_nums"] / self.error_dict["total_base_nums"])))
        self.get_txt(doc, "总体删除率：" + str(
            (self.error_dict["total_del_error_nums"] / self.error_dict["total_base_nums"])))
        self.get_txt(doc, "总体替换率：" + str(
            (self.error_dict["total_sub_error_nums"] / self.error_dict["total_base_nums"])))
        self.get_txt(doc, "总体增添率：" + str(
            (self.error_dict["total_ins_error_nums"] / self.error_dict["total_base_nums"])))
        self.get_txt(doc, "总错误序列率：" + str(
            (self.error_dict["total_error_reads_nums"] / self.error_dict["total_reads"])))
        self.get_txt(doc,
                     "总断链率：" + str((self.error_dict["total_broken_reads_nums"] / self.error_dict["total_reads"])))

        title_text = "总体比对情况"
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(8, 46, 84)
        # 添加比对情况
        txt3 = "比对成功率" + "缺失比对的ref数量" + "ref比对成功数分布图"
        self.get_txt(doc, "比对成功率:" + str(((self.error_dict["total_reads"] - self.error_dict[
            "total_no_align_nums"]) / self.error_dict["total_reads"])))
        self.get_txt(doc, "缺失比对的ref数量:" + str(self.error_dict["miss_ref_nums"]))
        self.get_txt(doc, "ref扩增数分布图")
        ref_nums_list = [0] * 300
        for key in self.error_dict["miss_ref_list"]:
            ref_nums_list.append(self.error_dict["miss_ref_list"][key])

        self.get_fig(doc, "ref扩增图", ref_nums_list, True)

        title_text = "碱基错误率情况"
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(8, 46, 84)
        # 添加碱基偏好性的影响：
        txt4 = "不同碱基的删除数" + "不同碱基的增添数" + "不同碱基的替换关联表格"
        self.get_txt(doc, "不同碱基的删除数：" + str(self.error_dict["base_del_nums"]))
        self.get_txt(doc, "不同碱基的增添数：" + str(self.error_dict["base_ins_nums"]))
        self.get_txt(doc, "不同碱基的替换关联表格：")
        self.get_table(doc, self.error_dict["base_sub_nums_matrix"])

        title_text = "错误率位置分布情况"
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(8, 46, 84)
        # 添加碱基位置影响：
        txt5 = "不同位置的碱基错误分布" + "不同位置的碱基删除错误分布" + "不同位置的碱基替换错误分布" + "不同位置的碱基增添错误分布"
        self.get_txt(doc, "不同位置的碱基错误分布：")
        self.get_fig(doc, "碱基错误分布", self.error_dict["loc_base_error_list"])
        self.get_txt(doc, "不同位置的碱基删除错误分布：")
        self.get_fig(doc, "删除碱基错误分布", self.error_dict["loc_base_del_error_list"])
        self.get_txt(doc, "不同位置的碱基替换错误分布：")
        self.get_fig(doc, "替换碱基错误分布", self.error_dict["loc_base_sub_error_list"])
        self.get_txt(doc, "不同位置的碱基增添错误分布：")
        self.get_fig(doc, "增添碱基错误分布", self.error_dict["loc_base_ins_error_list"])

        title_text = "错误关联性情况情况"
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(8, 46, 84)
        # 添加碱基关联性的情况
        txt6 = "前序删除关联矩阵" + "前序替换关联矩阵" + "前序增添关联矩阵" + "后序替换矩阵" + "后序增添矩阵"
        self.get_txt(doc, "前序删除关联矩阵：")
        self.get_table(doc, self.error_dict["pre_order_del_matrix"])
        self.get_txt(doc, "前序替换关联矩阵：")
        self.get_table(doc, self.error_dict["pre_order_sub_matrix"])
        self.get_txt(doc, "前序增添关联矩阵：")
        self.get_table(doc, self.error_dict["pre_order_ins_matrix"])
        self.get_txt(doc, "后序替换关联矩阵：")
        self.get_table(doc, self.error_dict["post_order_sub_error_matrix"])
        self.get_txt(doc, "后序增添关联矩阵：")
        self.get_table(doc, self.error_dict["post_order_ins_error_matrix"])

        title_text = "错误连续性情况"
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(8, 46, 84)
        # 添加碱基连续性错误情况
        txt7 = "连续性错误长度分布" + "连续性删除错误长度分布" + "连续性替换错误长度分布" + "连续性增添错误长度分布"
        self.get_txt(doc, "连续性错误长度分布：")
        self.get_fig(doc, "连续性错误", self.error_dict["contiguity_error_list"], False, True)
        self.get_txt(doc, "连续性删除错误长度分布：")
        self.get_fig(doc, "连续性删除错误", self.error_dict["contiguity_del_error_list"], False, True)
        self.get_txt(doc, "连续性替换错误长度分布：")
        self.get_fig(doc, "连续性替换错误", self.error_dict["contiguity_sub_error_list"], False, True)
        self.get_txt(doc, "连续性增添错误长度分布：")
        self.get_fig(doc, "连续性增添错误", self.error_dict["contiguity_ins_error_list"], False, True)

        title_text = "断链情况"
        title_paragraph = doc.add_paragraph(title_text)
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_paragraph.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(8, 46, 84)
        # 增加断链情况
        txt8 = "断链长度分布情况" + "断链位置关联矩阵"
        self.get_txt(doc, "断链长度分布：")
        self.get_fig(doc, "断链长度分布", self.error_dict["broken_error_list"], False, True)
        self.get_txt(doc, "断链位置关联矩阵：")
        self.get_table(doc, self.error_dict["broken_error_base_list"])

        output_name = self.name + ".docx"
        doc.save(output_name)

    def get_txt(self, doc, txt, title_mode=False):
        # 输入一段文字，把它放到word里面
        from docx.shared import Inches, Pt
        info = txt + '\n'
        p = doc.add_paragraph(info)
        if title_mode == True:
            p.style = 'Title'
        run = p.runs[0]
        font = run.font
        # font.name = '微软雅黑'
        font.size = Pt(12)  # 小四字体大小
        font.name = "MS Mincho"

    def get_table(self, doc, np_array):
        # 添加一个居中的段落作为表格
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml import OxmlElement
        table_paragraph = doc.add_paragraph()
        table_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 添加一个5x5的表格
        table = doc.add_table(rows=5, cols=5)
        table.style = 'Table Grid'

        # 设置第一行为 "0 A T G C"
        header_row = table.rows[0]
        header_row.cells[0].text = "0"
        header_row.cells[1].text = "A"
        header_row.cells[2].text = "T"
        header_row.cells[3].text = "G"
        header_row.cells[4].text = "C"

        # 填充第一列为 "0 A T G C"
        for i in range(1, 5):
            table.cell(i, 0).text = header_row.cells[i].text

        # 填充剩下的4x4部分为从0到16
        for i in range(1, 5):
            for j in range(1, 5):
                table.cell(i, j).text = str(np_array[i - 1, j - 1])
        doc.add_paragraph("\n")

    def get_fig(self, doc, name, input_list, frequency_mode=False, log_mode=False):
        import matplotlib.pyplot as plt
        import numpy as np
        from PIL import Image
        from docx import Document
        from docx.shared import Inches, Pt
        if frequency_mode == False:
            plt.figure(figsize=(10, 6))
            plt.bar(range(len(input_list)), input_list)
            if log_mode == True:
                plt.yscale('log')  # 设置纵坐标为对数坐标系
                plt.ylabel('Value (log scale)')
            else:
                plt.ylabel('Value')
            plt.xlabel('Index')

            plt.title('Bar Chart with Log Scale')
        else:
            from collections import Counter

            frequency_counter = Counter(input_list)
            # print(list,frequency_counter.keys())
            # 提取元素和频率
            elements = list(frequency_counter.keys())
            frequencies = list(frequency_counter.values())
            # 创建柱状图
            plt.figure(figsize=(10, 6))
            plt.bar(elements, frequencies)
            plt.xlabel('Elements')
            plt.ylabel('Frequency')
        fig_name = "uesless.jpg"
        plt.savefig(fig_name)
        doc.add_picture(fig_name, width=Inches(4))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = 1  # 设置为居中对齐
        doc.add_paragraph("\n")
        import os
        os.remove(fig_name)


if __name__ == '__main__':
    ref_path = "/mnt/st_data/linhaiyan/P10_5_BDDP210000009/ref.fasta"
    read_path = "/mnt/st_data/linhaiyan/P10_5_BDDP210000009/mem-se.sam"
    Error_model = Error_stats(ref_path, read_path)
    Error_model.process_read()








